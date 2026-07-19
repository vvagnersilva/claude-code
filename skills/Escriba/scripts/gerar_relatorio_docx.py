# /// script
# requires-python = ">=3.10"
# dependencies = ["python-docx"]
# ///
"""Monta um relatorio .docx a partir de um arquivo Markdown simples.

Uso:
    uv run gerar_relatorio_docx.py conteudo.md --saida relatorio.docx --titulo "Relatorio Mensal" --autor "Minha Empresa"

Markdown suportado (simples, suficiente para relatorios/laudos):
    # Titulo        -> Heading 1
    ## Secao        -> Heading 2
    ### Subsecao    -> Heading 3
    - item          -> lista com marcadores
    | a | b |       -> tabela (linhas consecutivas comecando com |)
    texto normal    -> paragrafo
"""
import argparse
import sys
from pathlib import Path

from docx import Document


def _add_tabela(doc, linhas_tabela: list[str]) -> None:
    linhas = []
    for ln in linhas_tabela:
        celulas = [c.strip() for c in ln.strip().strip("|").split("|")]
        if set("".join(celulas)) <= {"-", " ", ":"}:
            continue  # separador de cabecalho markdown
        linhas.append(celulas)
    if not linhas:
        return
    n_col = max(len(l) for l in linhas)
    tabela = doc.add_table(rows=0, cols=n_col)
    tabela.style = "Light Grid Accent 1"
    for l in linhas:
        celulas = tabela.add_row().cells
        for i in range(n_col):
            celulas[i].text = l[i] if i < len(l) else ""


def md_para_docx(texto_md: str, doc: Document) -> None:
    linhas = texto_md.splitlines()
    i = 0
    while i < len(linhas):
        ln = linhas[i]
        bruto = ln.rstrip()
        if bruto.startswith("### "):
            doc.add_heading(bruto[4:].strip(), level=3)
        elif bruto.startswith("## "):
            doc.add_heading(bruto[3:].strip(), level=2)
        elif bruto.startswith("# "):
            doc.add_heading(bruto[2:].strip(), level=1)
        elif bruto.startswith("- ") or bruto.startswith("* "):
            doc.add_paragraph(bruto[2:].strip(), style="List Bullet")
        elif bruto.startswith("|"):
            bloco = []
            while i < len(linhas) and linhas[i].lstrip().startswith("|"):
                bloco.append(linhas[i])
                i += 1
            _add_tabela(doc, bloco)
            continue
        elif bruto.strip() == "":
            pass
        else:
            doc.add_paragraph(bruto)
        i += 1


def main() -> int:
    p = argparse.ArgumentParser(description="Markdown -> relatorio DOCX.")
    p.add_argument("md", type=Path, help="Arquivo Markdown de entrada")
    p.add_argument("--saida", type=Path, default=None, help="Arquivo .docx de saida")
    p.add_argument("--titulo", default=None, help="Titulo no topo do documento")
    p.add_argument("--autor", default=None, help="Autor/empresa no cabecalho")
    args = p.parse_args()

    if not args.md.exists():
        print(f"Arquivo nao encontrado: {args.md}", file=sys.stderr)
        return 1

    doc = Document()
    if args.titulo:
        doc.add_heading(args.titulo, level=0)
    if args.autor:
        par = doc.add_paragraph(args.autor)
        par.runs[0].italic = True

    md_para_docx(args.md.read_text(encoding="utf-8"), doc)

    saida = args.saida or args.md.with_suffix(".docx")
    saida.parent.mkdir(parents=True, exist_ok=True)
    doc.save(saida)
    print(f"OK: relatorio salvo em {saida}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
